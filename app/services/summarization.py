import fitz  # PyMuPDF
from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
from transformers import LongformerTokenizer, LongformerForSequenceClassification
import os
import docx
import re
import boto3
from botocore.exceptions import NoCredentialsError

# Configure AWS credentials
aws_access_key_id = 'AKIAVRUVQANFUJNOWIUW'
aws_secret_access_key = '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'
aws_region = 'us-east-1'
bucket_name = "digimachine"

# Initialize S3 client
s3 = boto3.client(
    's3',
    aws_access_key_id=aws_access_key_id,
    aws_secret_access_key=aws_secret_access_key,
    region_name=aws_region
)

model_dir = "C:\\Users\\novneet.patnaik\\Documents\\GitHub\\ML-OCR-DM\\bart-large-cnn"

# Initialize the summarizer with a model and tokenizer from the local directory
tokenizer = AutoTokenizer.from_pretrained(model_dir)
model = AutoModelForSeq2SeqLM.from_pretrained(model_dir)
summarizer = pipeline("summarization", model=model, tokenizer=tokenizer, device=-1)


def extract_text_from_document(pdf_path):
    """
    Extract text from all pages of a PDF document.
    """
    try:
        document = fitz.open(pdf_path)
    except Exception as e:
        raise FileNotFoundError(f"Failed to open PDF file: {e}")

    if len(document) == 0:
        raise ValueError("The PDF file is empty.")

    text = ""
    for page_num in range(len(document)):
        page = document.load_page(page_num)
        text += page.get_text()

    if not text.strip():
        raise ValueError("The extracted text is empty.")

    return text


def extract_text_from_docx(docx_path):
    """
    Extract text from a DOCX document.
    """
    try:
        doc = docx.Document(docx_path)
    except Exception as e:
        raise FileNotFoundError(f"Failed to open DOCX file: {e}")

    if len(doc.paragraphs) == 0:
        raise ValueError("The DOCX file is empty.")

    text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])

    if not text.strip():
        raise ValueError("The extracted text is empty.")

    return text


def extract_text_from_txt(txt_path):
    """
    Extract text from a TXT document.
    """
    try:
        with open(txt_path, "r", encoding="utf-8") as file:
            text = file.read()
    except Exception as e:
        raise FileNotFoundError(f"Failed to open TXT file: {e}")

    if not text.strip():
        raise ValueError("The TXT file is empty.")

    return text


def identify_sections(text):
    """
    Identify and structure sections based on common heading patterns.
    """
    sections = {}
    current_section = "Introduction"
    sections[current_section] = ""

    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        # Detect headings based on common patterns (e.g., lines with all caps, numbered headings)
        if re.match(r'^[A-Z][A-Z\s]+$', line) or re.match(r'^\d+[\.\)]\s', line):
            current_section = line
            sections[current_section] = ""
        else:
            sections[current_section] += line + " "

    return sections


def chunk_text(text, tokenizer, max_length):
    """
    Splits the input text into chunks based on tokenization.
    """
    tokens = tokenizer.encode(text, truncation=True, max_length=max_length)
    chunks = []
    for i in range(0, len(tokens), max_length):
        chunks.append(tokens[i:i + max_length])
    return chunks


def summarize_section(section_text, summarizer, tokenizer, max_length=1024):
    """
    Summarize a section of text using the transformers pipeline.
    """
    chunks = chunk_text(section_text, tokenizer, max_length)
    chunk_summaries = [
        summarizer(tokenizer.decode(chunk), max_length=max_length, min_length=30, do_sample=False)[0]['summary_text']
        for chunk in chunks]
    return ' '.join(chunk_summaries)


def summarize_text(text, summarizer, tokenizer, max_length=1024):
    """
    Summarize the extracted text into structured sections.
    """
    sections = identify_sections(text)
    summary = {}

    for section, content in sections.items():
        summary[section] = summarize_section(content, summarizer, tokenizer, max_length)

    return summary


def save_summary_to_file(summary, output_file):
    """
    Save the summarized sections to a text file.
    """
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            for section, content in summary.items():
                f.write(f"{section}\n")
                f.write(f"{content.strip()}\n\n")
    except Exception as e:
        raise IOError(f"Failed to save summary to file: {e}")


def upload_to_s3(file_name, bucket, object_name=None):
    if object_name is None:
        object_name = os.path.basename(file_name)

    try:
        s3.upload_file(file_name, bucket, object_name)
    except NoCredentialsError:
        print("Credentials not available")
        return None

    return f"https://{bucket}.s3.amazonaws.com/{object_name}"


def summarize_document(file_path, output_file):
    """
    Main function to extract, summarize, and save the summary.
    """
    try:
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension == '.pdf':
            print("Extracting text from PDF...")
            text = extract_text_from_document(file_path)
        elif file_extension == '.docx':
            print("Extracting text from DOCX...")
            text = extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            print("Extracting text from TXT...")
            text = extract_text_from_txt(file_path)
        else:
            raise ValueError("Unsupported file format. Please upload a PDF, DOCX, or TXT file.")

        print("Summarizing text...")
        summary = summarize_text(text, summarizer, tokenizer, max_length=1024)

        print("Saving summary to file...")
        save_summary_to_file(summary, output_file)

        print("Uploading summary to S3...")
        download_link = upload_to_s3(output_file, bucket_name)

        if download_link:
            print(f"Summary saved successfully. Download link: {download_link}")
        else:
            print("Failed to upload summary to S3.")
    except FileNotFoundError as e:
        print(f"Error: {e}")
    except ValueError as e:
        print(f"Error: {e}")
    except RuntimeError as e:
        print(f"Error: {e}")
    except IOError as e:
        print(f"Error: {e}")

# if __name__ == "__main__":
#     file_path = "path/to/your/file.pdf"  # Update this with the path to your file (PDF, DOCX, or TXT)
#     output_file = "summary.txt"  # Output file for the summary

#     if not os.path.isfile(file_path):
#         print(f"Error: The file '{file_path}' does not exist.")
#     else:
#         main(file_path, output_file)
